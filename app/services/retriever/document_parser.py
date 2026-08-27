"""
retriever/document_parser.py

Extracts plain text from any supported file type.
Supported formats: PDF, DOCX, XLSX, XLS, TXT, CSV, MD, and any plain-text file.

Usage
-----
    parser = DocumentParser()
    chunks = parser.parse("/path/to/file.pdf")
    # Returns a list of text strings (one per page / sheet / logical chunk)
"""

from __future__ import annotations

import csv
import io
import os
from pathlib import Path


# ── Format handlers ───────────────────────────────────────────────────────────

def _parse_pdf(path: str) -> list[str]:
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        chunks = []
        for page in reader.pages:
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                chunks.append(text)
        return chunks
    except Exception as e:
        print(f"[DocumentParser] PDF error ({path}): {e}")
        return []


def _parse_docx(path: str) -> list[str]:
    try:
        import docx
        doc = docx.Document(path)
        # Group paragraphs into chunks of ~10 paragraphs each
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)
        # Group into chunks of 10 paragraphs
        chunks = []
        for i in range(0, len(paragraphs), 10):
            chunk = " ".join(paragraphs[i : i + 10])
            if chunk:
                chunks.append(chunk)
        return chunks
    except Exception as e:
        print(f"[DocumentParser] DOCX error ({path}): {e}")
        return []


def _parse_xlsx(path: str) -> list[str]:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        chunks = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join(
                    str(cell) for cell in row if cell is not None and str(cell).strip()
                )
                if row_text:
                    rows.append(row_text)
            if rows:
                # One chunk per sheet, split into groups of 20 rows
                for i in range(0, len(rows), 20):
                    chunk = f"[Sheet: {sheet_name}] " + " || ".join(rows[i : i + 20])
                    chunks.append(chunk)
        wb.close()
        return chunks
    except Exception as e:
        print(f"[DocumentParser] XLSX error ({path}): {e}")
        return []


def _parse_csv(path: str) -> list[str]:
    try:
        chunks = []
        with open(path, newline="", encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)
            rows = []
            for row in reader:
                row_text = " | ".join(cell.strip() for cell in row if cell.strip())
                if row_text:
                    rows.append(row_text)
        # Group rows into chunks of 30
        for i in range(0, len(rows), 30):
            chunks.append(" || ".join(rows[i : i + 30]))
        return chunks
    except Exception as e:
        print(f"[DocumentParser] CSV error ({path}): {e}")
        return []


def _parse_text(path: str) -> list[str]:
    """Header-preserving chunker for MD, TXT, and structured documents."""
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            lines = f.readlines()
        if not lines:
            return []

        chunks = []
        current_header = ""
        current_chunk_lines = []
        chunk_token_limit = 3000  # character limit per chunk to preserve tabular integrity

        for line in lines:
            stripped = line.strip()
            if stripped.startswith("#"):
                # When encountering a heading, finalize previous chunk if it reaches reasonable length
                if current_chunk_lines and sum(len(l) for l in current_chunk_lines) > 200:
                    chunk_text = (current_header + "\n" + "".join(current_chunk_lines)).strip()
                    chunks.append(chunk_text)
                    current_chunk_lines = []

                current_header = stripped

            current_chunk_lines.append(line)

            if sum(len(l) for l in current_chunk_lines) >= chunk_token_limit:
                chunk_text = (current_header + "\n" + "".join(current_chunk_lines)).strip()
                chunks.append(chunk_text)
                current_chunk_lines = []

        if current_chunk_lines:
            chunk_text = (current_header + "\n" + "".join(current_chunk_lines)).strip()
            chunks.append(chunk_text)

        return chunks
    except Exception as e:
        print(f"[DocumentParser] Text error ({path}): {e}")
        return []



# ── Extension → handler map ───────────────────────────────────────────────────

_HANDLERS = {
    ".pdf":  _parse_pdf,
    ".docx": _parse_docx,
    ".doc":  _parse_docx,   # python-docx handles .doc as well (best-effort)
    ".xlsx": _parse_xlsx,
    ".xls":  _parse_xlsx,
    ".csv":  _parse_csv,
    ".txt":  _parse_text,
    ".md":   _parse_text,
    ".rst":  _parse_text,
    ".json": _parse_text,
    ".html": _parse_text,
    ".htm":  _parse_text,
}

SUPPORTED_EXTENSIONS = set(_HANDLERS.keys())


# ── Public class ──────────────────────────────────────────────────────────────

class DocumentParser:
    """
    Parses a file into a list of plain-text chunks.

    Each chunk is a string of a few hundred to a few thousand characters —
    small enough to embed meaningfully, large enough to carry context.

    Usage
    -----
        parser = DocumentParser()
        chunks = parser.parse("report.pdf")
        # → ["Page 1 text...", "Page 2 text...", ...]
    """

    def parse(self, path: str) -> list[str]:
        """
        Parse *path* into a list of text chunks.
        Returns an empty list if the file is unsupported or unreadable.
        """
        ext = Path(path).suffix.lower()
        handler = _HANDLERS.get(ext, _parse_text)
        chunks = handler(path)
        # Attach source metadata to each chunk so retrieval results are traceable
        filename = os.path.basename(path)
        tagged = [f"[Source: {filename}] {chunk}" for chunk in chunks if chunk.strip()]
        return tagged

    def is_supported(self, path: str) -> bool:
        """Return True if this file type can be parsed."""
        return Path(path).suffix.lower() in SUPPORTED_EXTENSIONS